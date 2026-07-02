"""
Resume Parser Service
=====================
Handles text extraction from PDF and DOCX/DOC resume files, with:
  - DOCX: full document-order extraction including table cells (BUG 1 fix)
  - PDF:  pdfplumber direct extraction + PyMuPDF OCR fallback for
          vector-only / image-only PDFs (BUG 2 fix)
  - Structured logging whenever extraction yields suspiciously little text
    or falls back to OCR (GENERAL fix)
"""
from __future__ import annotations

import logging
import os
from io import BytesIO
from typing import Optional, Tuple

import pdfplumber
from docx import Document
from docx.oxml.ns import qn

from app.models.schemas import ResumeData, ExperienceItem, EducationItem, ContactInfo
from app.services.openai_client import get_openai_service

# --------------------------------------------------------------------------- #
# Logger – emit structured key=value pairs so log aggregators can parse them  #
# --------------------------------------------------------------------------- #
logger = logging.getLogger(__name__)

# Threshold: if extracted text is shorter than this we flag it as suspicious.
_THIN_TEXT_THRESHOLD = 300

# Minimum number of curves/rects/images on a page to consider it "has visual
# content" and therefore worth attempting OCR (avoids OCR on blank pages).
_OCR_VISUAL_ELEMENT_MIN = 5


# --------------------------------------------------------------------------- #
# Standalone extraction helpers (also importable for unit-testing)             #
# --------------------------------------------------------------------------- #

def _cell_text(cell) -> str:
    """Return all text from a DOCX table cell, joining inner paragraphs."""
    return " ".join(
        p.text for p in cell.paragraphs if p.text.strip()
    )


def _table_text(table) -> str:
    """
    Extract text from a DOCX table (including nested tables) row-by-row,
    cell-by-cell.  Cells are separated by ' | ', rows by newlines.
    Nested tables are resolved recursively before joining with the
    surrounding cell text.
    """
    rows_text: list[str] = []
    for row in table.rows:
        cells_text: list[str] = []
        for cell in row.cells:
            # Recurse into any nested tables inside this cell
            nested_parts: list[str] = []
            for nested_tbl in cell.tables:
                nested_parts.append(_table_text(nested_tbl))
            # Direct paragraph text of this cell
            direct = _cell_text(cell)
            if nested_parts:
                combined = (direct + "\n" + "\n".join(nested_parts)).strip()
            else:
                combined = direct
            if combined:
                cells_text.append(combined)
        if cells_text:
            rows_text.append(" | ".join(cells_text))
    return "\n".join(rows_text)


def extract_docx_text(file_content: bytes) -> str:
    """
    Extract text from a DOCX (or DOC via python-docx) file in document order.

    Walks the XML body element-by-element so that paragraphs and tables appear
    in the same reading order as the original document.  Table cells are joined
    with ' | ' (columns) and '\\n' (rows) so structure is preserved.
    Nested tables are handled recursively.

    This replaces the old `doc.paragraphs`-only loop which silently dropped
    all text inside tables.
    """
    doc = Document(BytesIO(file_content))
    body = doc.element.body
    parts: list[str] = []

    for child in body:
        tag = child.tag

        # w:p  →  plain paragraph
        if tag == qn("w:p"):
            # Use the python-docx Paragraph API for clean .text access
            from docx.text.paragraph import Paragraph as DocxParagraph
            para = DocxParagraph(child, doc)
            text = para.text.strip()
            if text:
                parts.append(text)

        # w:tbl  →  table (may contain nested tables)
        elif tag == qn("w:tbl"):
            from docx.table import Table as DocxTable
            tbl = DocxTable(child, doc)
            tbl_text = _table_text(tbl)
            if tbl_text.strip():
                parts.append(tbl_text)

        # Any other block-level element: skip (sectPr, bookmarks, etc.)

    return "\n".join(parts)


def extract_pdf_text(file_content: bytes, filename: str = "resume.pdf") -> str:
    """
    Extract text from a PDF file.

    Strategy:
      1. Try pdfplumber's direct text extraction (fast, accurate for
         text-layer PDFs).
      2. For pages where direct extraction yields nothing but there is clearly
         visual content (curves/rects/images), fall back to PyMuPDF's built-in
         text extraction on a rasterised page image (OCR-equivalent).

    Logs clearly whenever a page or whole document falls into the OCR path.
    """
    direct_parts: list[str] = []
    ocr_parts: list[str] = []
    pages_ocr_attempted = 0
    pages_ocr_succeeded = 0

    with pdfplumber.open(BytesIO(file_content)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""

            if text.strip():
                direct_parts.append(text)
                continue

            # --- No text layer on this page ---
            visual_elements = (
                len(page.curves or [])
                + len(page.rects or [])
                + len(page.images or [])
            )

            if visual_elements < _OCR_VISUAL_ELEMENT_MIN:
                # Genuinely blank page – skip silently.
                logger.debug(
                    "resume_parser: page=%d file=%s is blank (no text, few visuals=%d), skipping",
                    page_num, filename, visual_elements,
                )
                continue

            # Visual content present but no text → needs OCR fallback
            pages_ocr_attempted += 1
            logger.warning(
                "resume_parser: OCR_FALLBACK_TRIGGERED file=%s page=%d "
                "visual_elements=%d – pdfplumber returned no text",
                filename, page_num, visual_elements,
            )

            ocr_text = _ocr_page_via_pymupdf(file_content, page_num - 1, filename)
            if ocr_text.strip():
                ocr_parts.append(ocr_text)
                pages_ocr_succeeded += 1
            else:
                logger.error(
                    "resume_parser: OCR_FAILED file=%s page=%d – "
                    "PyMuPDF returned no text after rasterisation",
                    filename, page_num,
                )

    if pages_ocr_attempted:
        logger.info(
            "resume_parser: OCR_SUMMARY file=%s pages_ocr_attempted=%d "
            "pages_ocr_succeeded=%d",
            filename, pages_ocr_attempted, pages_ocr_succeeded,
        )

    return "\n".join(direct_parts + ocr_parts)


def _ocr_page_via_pymupdf(file_content: bytes, page_index: int, filename: str) -> str:
    """
    Rasterise a single PDF page with PyMuPDF and use its built-in
    Tesseract-free text layer extraction on the pixmap.

    PyMuPDF (fitz) can render a page to a pixmap and then run its own
    textpage extraction from the rendered image using `get_text("text")`
    after re-opening with the rendered bytes – but more practically we use
    `page.get_text("text")` directly via fitz, which reads font-based text
    even from pages pdfplumber misses (different PDF parser underneath).

    If fitz also returns nothing (true image-only PDF), we log and return "".
    """
    try:
        import fitz  # PyMuPDF
        
        # Ensure Tesseract can be found by PyMuPDF on Windows
        import os
        if os.name == 'nt' and 'TESSDATA_PREFIX' not in os.environ:
            tessdata_path = r"C:\Program Files\Tesseract-OCR\tessdata"
            if os.path.exists(tessdata_path):
                os.environ['TESSDATA_PREFIX'] = tessdata_path

    except ImportError:
        logger.error(
            "resume_parser: PyMuPDF (fitz) not installed – cannot OCR page %d of %s",
            page_index + 1, filename,
        )
        return ""

    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        page = doc[page_index]

        # First: try fitz's own text extraction (catches more PDFs than pdfplumber)
        text = page.get_text("text").strip()
        if text:
            logger.info(
                "resume_parser: OCR_PYMUPDF_TEXT_SUCCESS file=%s page=%d chars=%d",
                filename, page_index + 1, len(text),
            )
            doc.close()
            return text

        # Second: rasterise at 300 DPI and try fitz's OCR layer if available
        # (requires Tesseract installed system-wide + PyMuPDF built with OCR support)
        try:
            mat = fitz.Matrix(300 / 72, 300 / 72)   # 300 DPI scale matrix
            pix = page.get_pixmap(matrix=mat, alpha=False)
            # get_textpage_ocr is available in PyMuPDF >= 1.21
            tp = page.get_textpage_ocr(flags=0, dpi=300, full=True)
            ocr_text = tp.extractText().strip() if tp else ""
            if ocr_text:
                logger.info(
                    "resume_parser: OCR_PYMUPDF_OCR_SUCCESS file=%s page=%d chars=%d",
                    filename, page_index + 1, len(ocr_text),
                )
                doc.close()
                return ocr_text
        except (AttributeError, Exception) as ocr_err:
            # get_textpage_ocr not available or Tesseract not installed
            logger.warning(
                "resume_parser: OCR_PYMUPDF_OCR_UNAVAILABLE file=%s page=%d reason=%s",
                filename, page_index + 1, ocr_err,
            )

        doc.close()
        return ""

    except Exception as exc:
        logger.error(
            "resume_parser: OCR_EXCEPTION file=%s page=%d error=%s",
            filename, page_index + 1, exc,
        )
        return ""


# --------------------------------------------------------------------------- #
# Service class                                                                 #
# --------------------------------------------------------------------------- #

class ResumeParserService:
    """Service for parsing resumes using AI-powered semantic extraction."""

    def __init__(self):
        self.openai = get_openai_service()

    async def parse_resume(self, file_content: bytes, filename: str) -> Tuple[str, ResumeData]:
        """
        Parse a resume file and extract structured data.
        Returns tuple of (raw_text, parsed_data).
        """
        raw_text = await self._extract_text(file_content, filename)
        parsed_data = await self._ai_parse_resume(raw_text)
        return raw_text, parsed_data

    async def _extract_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from PDF or DOCX file with structured logging."""
        ext = os.path.splitext(filename)[1].lower()

        if ext == ".pdf":
            method = "pdfplumber+pymupdf_ocr_fallback"
            text = extract_pdf_text(file_content, filename)
            
            if not text.strip():
                try:
                    import fitz
                    import base64
                    doc = fitz.open(stream=file_content, filetype="pdf")
                    base64_images = []
                    # Process at most the first 3 pages to save tokens
                    for i in range(min(3, len(doc))):
                        page = doc[i]
                        mat = fitz.Matrix(2.0, 2.0) # approx 144 DPI
                        pix = page.get_pixmap(matrix=mat, alpha=False)
                        img_b64 = base64.b64encode(pix.tobytes("jpeg")).decode("utf-8")
                        base64_images.append(img_b64)
                    doc.close()
                    
                    if base64_images:
                        logger.warning("resume_parser: VISUAL_OCR_FALLBACK triggering OpenAI Vision for %s", filename)
                        method = "openai_vision_ocr"
                        text = await self.openai.extract_text_from_images(base64_images)
                except Exception as e:
                    logger.error("resume_parser: VISION_OCR_FAILED file=%s error=%s", filename, e)

        elif ext in [".docx", ".doc"]:
            method = "python_docx_full_body"
            text = extract_docx_text(file_content)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        char_count = len(text)

        if char_count == 0:
            logger.error(
                "resume_parser: EXTRACTION_EMPTY file=%s ext=%s method=%s char_count=0 "
                "– ATS screening will be skipped for this resume",
                filename, ext, method,
            )
        elif char_count < _THIN_TEXT_THRESHOLD:
            logger.warning(
                "resume_parser: EXTRACTION_THIN file=%s ext=%s method=%s char_count=%d "
                "(threshold=%d) – ATS score may be unreliable",
                filename, ext, method, char_count, _THIN_TEXT_THRESHOLD,
            )
        else:
            logger.info(
                "resume_parser: EXTRACTION_OK file=%s ext=%s method=%s char_count=%d",
                filename, ext, method, char_count,
            )

        return text

    async def _ai_parse_resume(self, resume_text: str) -> ResumeData:
        """Use AI to semantically parse resume content."""
        try:
            result = await self.openai.analyze_resume(resume_text[:8000])

            experience = []
            for exp in result.get("experience", []):
                experience.append(ExperienceItem(
                    title=exp.get("title", ""),
                    company=exp.get("company", ""),
                    duration=exp.get("duration", ""),
                    description=exp.get("description", ""),
                    start_date=exp.get("start_date"),
                    end_date=exp.get("end_date"),
                ))

            education = []
            for edu in result.get("education", []):
                education.append(EducationItem(
                    degree=edu.get("degree", ""),
                    institution=edu.get("institution", ""),
                    year=edu.get("year", ""),
                ))

            contact_data = result.get("contact", {})
            contact = ContactInfo(
                email=contact_data.get("email"),
                phone=contact_data.get("phone"),
                linkedin=contact_data.get("linkedin"),
            )

            return ResumeData(
                skills=result.get("skills", []),
                experience=experience,
                education=education,
                summary=result.get("summary") or "",
                contact=contact,
                total_experience_years=float(result.get("total_experience_years", 0)),
                certifications=result.get("certifications", []),
            )

        except Exception as e:
            logger.error("resume_parser: AI_PARSE_FAILED error=%s", e)
            return ResumeData()


# --------------------------------------------------------------------------- #
# Singleton                                                                     #
# --------------------------------------------------------------------------- #

_resume_parser: Optional[ResumeParserService] = None


def get_resume_parser() -> ResumeParserService:
    global _resume_parser
    if _resume_parser is None:
        _resume_parser = ResumeParserService()
    return _resume_parser
