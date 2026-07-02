"""
Tests for resume_parser.py — BUG 1 (DOCX table extraction) and BUG 2 (OCR logging).

Run from backend/ directory:
    pytest tests/test_resume_parser.py -v
"""
from __future__ import annotations

import io
import logging
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --------------------------------------------------------------------------- #
# Helper: build a synthetic DOCX in memory                                     #
# --------------------------------------------------------------------------- #

def _make_two_column_resume_docx() -> bytes:
    """
    Creates a DOCX with:
      - One plain paragraph before the table ("John Doe – Software Engineer")
      - A 4-row, 2-column table with resume-style content:
          Summary    | Experienced developer with 5 years in Python and cloud.
          Skills     | Python | FastAPI | Docker | AWS
          Experience | Senior Developer at TechCorp (2020-2024)
          Education  | B.Tech Computer Science, XYZ University 2018
    Returns the document as raw bytes.
    """
    doc = Document()

    # Plain paragraph BEFORE the table
    doc.add_paragraph("John Doe – Software Engineer")

    # 2-column table (label | content)
    table = doc.add_table(rows=4, cols=2)

    rows_data = [
        ("Summary",    "Experienced developer with 5 years in Python and cloud."),
        ("Skills",     "Python | FastAPI | Docker | AWS"),
        ("Experience", "Senior Developer at TechCorp (2020-2024)"),
        ("Education",  "B.Tech Computer Science, XYZ University 2018"),
    ]
    for row_idx, (label, content) in enumerate(rows_data):
        table.rows[row_idx].cells[0].text = label
        table.rows[row_idx].cells[1].text = content

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_nested_table_docx() -> bytes:
    """
    Creates a DOCX where one cell contains a nested table.
    Outer table: 1 row, 1 col whose cell contains a nested 2x1 table.
    """
    doc = Document()
    doc.add_paragraph("Nested table test")

    outer = doc.add_table(rows=1, cols=1)
    outer_cell = outer.rows[0].cells[0]

    # Add nested table into the outer cell
    nested = outer_cell.add_table(rows=2, cols=1)
    nested.rows[0].cells[0].text = "Nested Row 1"
    nested.rows[1].cells[0].text = "Nested Row 2"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# BUG 1 — DOCX table extraction                                                #
# --------------------------------------------------------------------------- #

class TestExtractDocxText:
    """Tests for extract_docx_text() – BUG 1 fix."""

    def test_paragraph_before_table_is_included(self):
        from app.services.resume_parser import extract_docx_text

        docx_bytes = _make_two_column_resume_docx()
        result = extract_docx_text(docx_bytes)

        assert "John Doe" in result, "Header paragraph should be in output"

    def test_all_table_cell_labels_extracted(self):
        from app.services.resume_parser import extract_docx_text

        docx_bytes = _make_two_column_resume_docx()
        result = extract_docx_text(docx_bytes)

        for label in ("Summary", "Skills", "Experience", "Education"):
            assert label in result, f"Table label '{label}' should appear in extracted text"

    def test_all_table_cell_content_extracted(self):
        from app.services.resume_parser import extract_docx_text

        docx_bytes = _make_two_column_resume_docx()
        result = extract_docx_text(docx_bytes)

        expected_snippets = [
            "Experienced developer with 5 years in Python and cloud.",
            "Python",
            "FastAPI",
            "Docker",
            "AWS",
            "Senior Developer at TechCorp",
            "B.Tech Computer Science",
            "XYZ University",
        ]
        for snippet in expected_snippets:
            assert snippet in result, f"Expected snippet not found in extracted text: '{snippet}'"

    def test_document_order_respected(self):
        """Header paragraph should appear before table content."""
        from app.services.resume_parser import extract_docx_text

        docx_bytes = _make_two_column_resume_docx()
        result = extract_docx_text(docx_bytes)

        header_pos = result.find("John Doe")
        summary_pos = result.find("Summary")
        assert header_pos < summary_pos, (
            "Header paragraph should appear before table rows in document order"
        )

    def test_nested_table_content_extracted(self):
        """Nested tables inside cells must be recursed into."""
        from app.services.resume_parser import extract_docx_text

        docx_bytes = _make_nested_table_docx()
        result = extract_docx_text(docx_bytes)

        assert "Nested Row 1" in result, "Nested table row 1 should be extracted"
        assert "Nested Row 2" in result, "Nested table row 2 should be extracted"

    def test_old_paragraphs_only_approach_misses_table(self):
        """
        Regression guard: confirm the OLD approach (doc.paragraphs only)
        would have missed the table content, proving our fix is necessary.
        """
        from docx import Document as DocxDocument

        docx_bytes = _make_two_column_resume_docx()
        doc = DocxDocument(io.BytesIO(docx_bytes))
        old_result = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        # The old approach should NOT have 'Summary' from the table
        assert "Summary" not in old_result, (
            "The OLD approach (doc.paragraphs) should miss table content – "
            "if this assertion fails the test itself is broken"
        )

    def test_empty_docx_returns_empty_string(self):
        """A DOCX with no content should not crash and return empty string."""
        from app.services.resume_parser import extract_docx_text

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        result = extract_docx_text(buf.getvalue())
        assert result == ""


# --------------------------------------------------------------------------- #
# BUG 2 — PDF OCR fallback logging                                             #
# --------------------------------------------------------------------------- #

class TestExtractPdfText:
    """Tests for extract_pdf_text() – BUG 2 fix (logging & OCR path)."""

    def test_normal_pdf_uses_direct_extraction(self, caplog):
        """A PDF with a real text layer should never trigger OCR."""
        import pdfplumber

        # Build a minimal real-text PDF using reportlab if available,
        # otherwise skip (we test the logic path, not PDF generation)
        try:
            from reportlab.pdfgen import canvas as rl_canvas
            from reportlab.lib.pagesizes import A4

            buf = io.BytesIO()
            c = rl_canvas.Canvas(buf, pagesize=A4)
            c.drawString(72, 720, "Hello World Resume")
            c.save()
            pdf_bytes = buf.getvalue()
        except ImportError:
            pytest.skip("reportlab not installed – skipping PDF text-layer test")

        from app.services.resume_parser import extract_pdf_text

        with caplog.at_level(logging.WARNING):
            result = extract_pdf_text(pdf_bytes, "test_normal.pdf")

        # OCR fallback should NOT have been triggered
        assert "OCR_FALLBACK_TRIGGERED" not in caplog.text
        # Some text should be present
        assert len(result) > 0

    def test_empty_page_with_visuals_triggers_ocr_log(self, caplog):
        """
        Simulate a page where pdfplumber returns no text but there are curves.
        Verify OCR_FALLBACK_TRIGGERED is logged.
        """
        from app.services.resume_parser import extract_pdf_text

        # Mock pdfplumber to return a page with no text but >5 curves
        mock_page = MagicMock()
        mock_page.extract_text.return_value = ""
        mock_page.curves = [{}] * 10  # 10 curves → visual content
        mock_page.rects = []
        mock_page.images = []

        mock_pdf = MagicMock()
        mock_pdf.pages = [mock_page]
        mock_pdf.__enter__ = lambda s: mock_pdf
        mock_pdf.__exit__ = MagicMock(return_value=False)

        fake_bytes = b"%PDF-1.4 fake"

        with patch("pdfplumber.open", return_value=mock_pdf), \
             patch("app.services.resume_parser._ocr_page_via_pymupdf", return_value="OCR text here"), \
             caplog.at_level(logging.WARNING):
            result = extract_pdf_text(fake_bytes, "vector_resume.pdf")

        assert "OCR_FALLBACK_TRIGGERED" in caplog.text
        assert "vector_resume.pdf" in caplog.text
        assert "OCR text here" in result

    def test_blank_page_no_visuals_does_not_trigger_ocr(self, caplog):
        """A truly blank page (no text, no visuals) should be silently skipped."""
        from app.services.resume_parser import extract_pdf_text

        mock_page = MagicMock()
        mock_page.extract_text.return_value = ""
        mock_page.curves = []
        mock_page.rects = []
        mock_page.images = []

        mock_pdf = MagicMock()
        mock_pdf.pages = [mock_page]
        mock_pdf.__enter__ = lambda s: mock_pdf
        mock_pdf.__exit__ = MagicMock(return_value=False)

        with patch("pdfplumber.open", return_value=mock_pdf), \
             caplog.at_level(logging.WARNING):
            result = extract_pdf_text(b"%PDF fake", "blank.pdf")

        assert "OCR_FALLBACK_TRIGGERED" not in caplog.text
        assert result == ""

    def test_ocr_failure_is_logged_as_error(self, caplog):
        """If OCR is triggered but returns nothing, an error is logged."""
        from app.services.resume_parser import extract_pdf_text

        mock_page = MagicMock()
        mock_page.extract_text.return_value = ""
        mock_page.curves = [{}] * 10
        mock_page.rects = []
        mock_page.images = []

        mock_pdf = MagicMock()
        mock_pdf.pages = [mock_page]
        mock_pdf.__enter__ = lambda s: mock_pdf
        mock_pdf.__exit__ = MagicMock(return_value=False)

        with patch("pdfplumber.open", return_value=mock_pdf), \
             patch("app.services.resume_parser._ocr_page_via_pymupdf", return_value=""), \
             caplog.at_level(logging.ERROR):
            result = extract_pdf_text(b"%PDF fake", "image_only.pdf")

        assert "OCR_FAILED" in caplog.text
        assert result == ""


# --------------------------------------------------------------------------- #
# GENERAL — structured logging in _extract_text                                #
# --------------------------------------------------------------------------- #

class TestExtractionLogging:
    """Tests for structured logging in ResumeParserService._extract_text."""

    @pytest.mark.asyncio
    async def test_thin_text_warns(self, caplog):
        """If extracted text is < 300 chars, a WARNING is emitted."""
        from app.services.resume_parser import ResumeParserService

        # Build a real DOCX with very little content
        doc = Document()
        doc.add_paragraph("Hi")
        buf = io.BytesIO()
        doc.save(buf)
        tiny_docx = buf.getvalue()

        service = ResumeParserService.__new__(ResumeParserService)

        with caplog.at_level(logging.WARNING, logger="app.services.resume_parser"):
            await service._extract_text(tiny_docx, "tiny.docx")

        assert "EXTRACTION_THIN" in caplog.text or "EXTRACTION_EMPTY" in caplog.text

    @pytest.mark.asyncio
    async def test_good_extraction_logs_info(self, caplog):
        """A well-extracted resume logs EXTRACTION_OK at INFO level."""
        from app.services.resume_parser import ResumeParserService

        # Build a DOCX with sufficient content
        doc = Document()
        doc.add_paragraph("A" * 400)
        buf = io.BytesIO()
        doc.save(buf)
        rich_docx = buf.getvalue()

        service = ResumeParserService.__new__(ResumeParserService)

        with caplog.at_level(logging.INFO, logger="app.services.resume_parser"):
            await service._extract_text(rich_docx, "rich.docx")

        assert "EXTRACTION_OK" in caplog.text

    @pytest.mark.asyncio
    async def test_unsupported_extension_raises(self):
        """An unsupported extension should raise ValueError immediately."""
        from app.services.resume_parser import ResumeParserService

        service = ResumeParserService.__new__(ResumeParserService)

        with pytest.raises(ValueError, match="Unsupported file format"):
            await service._extract_text(b"data", "resume.rtf")
